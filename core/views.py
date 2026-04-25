from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, Q
from .models import Produto, Projeto, Empresa, Perfil, ItemProjeto, ItemMO
from functools import wraps


# ============================================================
# HELPERS DE CARGO / PERMISSÃO
# ============================================================

# Grupos de cargo para controle de acesso
CARGOS_DIRETORIA = {'diretor_geral', 'diretor_setor'}
CARGOS_GERENCIA  = {'gerente_pre_vendas', 'supervisor'}
CARGOS_COMPRAS   = {'gerente_compras', 'comprador', 'orcamentista'}
CARGOS_ADMIN     = CARGOS_DIRETORIA  # quem pode gerir usuários/empresas/programa

def get_cargo(user):
    """Retorna o cargo do usuário ou 'analista' como fallback."""
    try:
        return user.perfil.cargo
    except Exception:
        return 'analista'

def tem_acesso(user, *cargos_permitidos):
    """True se o cargo do usuário estiver na lista de permitidos."""
    return get_cargo(user) in set(cargos_permitidos)

def cargo_required(*cargos_permitidos):
    """Decorator que bloqueia acesso se cargo não estiver na lista."""
    def decorator(view_func):
        @wraps(view_func)
        @login_required(login_url='/')
        def wrapper(request, *args, **kwargs):
            if get_cargo(request.user) not in set(cargos_permitidos):
                messages.error(request, 'Você não tem permissão para acessar esta página.')
                return redirect('inicio')
            return view_func(request, *args, **kwargs)
        return wrapper
    return decorator

def dash_url_para_cargo(cargo):
    """Retorna a URL do dashboard adequado para o cargo."""
    if cargo in CARGOS_COMPRAS:
        return 'dash_orcamentista'
    if cargo == 'analista':
        return 'dash_analista'
    # diretores, gerentes pré-vendas, supervisor, comercial, executivo
    return 'dash_gerencial'


FAIXAS_CATEGORIA = {
    'CFTV': 100000, 'VMS': 200000, 'ALARME': 300000, 'ACESSO': 400000,
    'CABEAMENTO': 500000, 'INFRA': 600000, 'REDE': 700000, 'SWITCHES': 750000,
    'ENERGIA': 800000, 'OUTROS': 900000, 'SERVICOS': 1000000,
}

def gerar_proximo_id(categoria):
    categoria_upper = categoria.upper().strip()
    mapa = {'ACESSO (C.A)': 'ACESSO', 'ELETRICIDADE': 'ENERGIA', 'SERVIÇO': 'SERVICOS', 'SERVICO': 'SERVICOS'}
    chave = mapa.get(categoria_upper, categoria_upper)
    faixa_inicio = FAIXAS_CATEGORIA.get(chave, 900000)
    ultimo = Produto.objects.filter(id_planilha__gte=faixa_inicio, id_planilha__lte=faixa_inicio+9999).order_by('-id_planilha').first()
    return (ultimo.id_planilha + 1) if ultimo else (faixa_inicio + 1)


# ============================================================
# LOGIN
# ============================================================
def minha_tela_login(request):
    if request.user.is_authenticated:
        return redirect('inicio')
    if request.method == "POST":
        login_input = request.POST.get('username')
        senha_digitada = request.POST.get('password')
        empresa_id = request.POST.get('company', '')
        username_final = login_input
        if "@" in login_input:
            try:
                username_final = User.objects.get(email=login_input).username
            except User.DoesNotExist:
                pass
        user = authenticate(request, username=username_final, password=senha_digitada)
        if user is not None:
            # Verificar se o usuário pertence à empresa selecionada
            if empresa_id and not user.is_superuser:
                try:
                    perfil_usr = user.perfil
                    if perfil_usr.empresa_id and str(perfil_usr.empresa_id) != str(empresa_id):
                        messages.error(request, "Usuário não autorizado para esta empresa.")
                        return render(request, 'login.html', {'empresas': Empresa.objects.all().order_by('nome_fantasia')})
                except Exception:
                    pass
            login(request, user)
            if empresa_id:
                request.session['empresa_id'] = int(empresa_id)
                request.session.modified = True
                request.session.save()
            return redirect('inicio')
        else:
            messages.error(request, "Usuário/E-mail ou senha incorretos!")
    empresas = Empresa.objects.all().order_by('nome_fantasia')
    return render(request, 'login.html', {'empresas': empresas})


def deslogar(request):
    logout(request)
    return redirect('login')


# ============================================================
# INÍCIO (tela principal)
# ============================================================
@login_required(login_url='/')
def inicio(request):
    total_projetos = Projeto.objects.filter(finalizado=True).count()
    total_produtos = Produto.objects.count()
    # Soma real dos projetos finalizados
    from decimal import Decimal
    valor_total_projetos = Decimal('0')
    for projeto in Projeto.objects.filter(finalizado=True):
        for item in ItemProjeto.objects.filter(projeto=projeto):
            valor_total_projetos += item.preco_unitario * item.quantidade
    # Formatar
    if valor_total_projetos >= 1000000:
        valor_fmt = f"{float(valor_total_projetos)/1000000:,.1f}M".replace(',', '.')
    elif valor_total_projetos >= 1000:
        valor_fmt = f"{float(valor_total_projetos):,.2f}".replace(',', 'X').replace('.', '.').replace('X', ',')
    else:
        valor_fmt = f"{float(valor_total_projetos):,.2f}".replace(',', 'X').replace('.', '.').replace('X', ',')
    context = {
        'total_projetos': total_projetos,
        'total_produtos': total_produtos,
        'valor_total_projetos': valor_fmt,
    }
    return render(request, 'inicio.html', context)


# ============================================================
# PRÉ-PROJETO (wizard etapa 1)
# ============================================================
@login_required(login_url='/')
def novo_projeto(request):
    if request.method == 'POST':
        id_manual = request.POST.get('id_projeto_manual', '').strip()
        if not id_manual:
            messages.error(request, 'O ID do projeto é obrigatório.')
            return render(request, 'novo_projeto.html')
        if Projeto.objects.filter(id_projeto_manual=id_manual).exists():
            messages.error(request, f'Já existe um projeto com o número "{id_manual}". Escolha outro.')
            return render(request, 'novo_projeto.html')
        try:
            projeto = Projeto.objects.create(
                id_projeto_manual=id_manual,
                nome_cliente=request.POST.get('nome_cliente', ''),
                estado_obra=request.POST.get('estado_obra', ''),
                municipio_obra=request.POST.get('municipio_obra', ''),
                empresa_executora=request.POST.get('empresa_executora', 'EMIVE'),
                faturamento_servico=request.POST.get('faturamento_servico', 'MENSAL'),
                tipo_projeto=request.POST.get('tipo_projeto', 'Privado'),
                licitacao_publicada=request.POST.get('licitacao_publicada') == 'true',
                finalizado=False,
            )
            return redirect('bom_selector_projeto', projeto_id=projeto.id)
        except Exception as e:
            messages.error(request, f'Erro ao criar pré-projeto: {e}')
    
    # Pega empresa da sessão para pré-selecionar
    empresa_logada = None
    try:
        empresa_id = request.session.get('empresa_id')
        if empresa_id:
            empresa_logada = Empresa.objects.get(id=empresa_id).nome_fantasia
    except Exception:
        pass
    return render(request, 'novo_projeto.html', {'empresa_logada': empresa_logada})


# ============================================================
# CONSULTAR PROJETO (mostra todos se sem filtro)
# ============================================================
@login_required(login_url='/')
def consultar_projeto(request):
    # Excluir projeto
    if request.method == 'POST':
        projeto_id = request.POST.get('excluir_projeto_id')
        if projeto_id:
            try:
                projeto = Projeto.objects.get(id=projeto_id)
                ItemProjeto.objects.filter(projeto=projeto).delete()
                projeto.delete()
            except Projeto.DoesNotExist:
                pass
        return redirect('consultar_projeto')

    q = request.GET.get('q', '').strip()
    if q:
        projetos = Projeto.objects.filter(
            Q(id_projeto_manual__icontains=q) | Q(nome_cliente__icontains=q)
        ).order_by('-data_criacao')
    else:
        projetos = Projeto.objects.all().order_by('-data_criacao')

    # Calcular totais para cada projeto
    projetos_dados = []
    for p in projetos:
        itens = ItemProjeto.objects.filter(projeto=p)
        custo = sum(float(i.preco_unitario) * i.quantidade for i in itens)
        tem_itens = itens.exists()
        projetos_dados.append({
            'projeto': p,
            'custo': custo,
            'venda': custo,  # TODO: aplicar MKP
            'tem_itens': tem_itens,
        })

    return render(request, 'consultar_projeto.html', {'projetos_dados': projetos_dados, 'query': q, 'total': len(projetos_dados)})


# ============================================================
# BOM SELECTOR
# ============================================================
@login_required(login_url='/')
def bom_selector(request, projeto_id=None):
    projeto = get_object_or_404(Projeto, id=projeto_id) if projeto_id else None

    # POST = incluir itens no projeto
    if request.method == 'POST' and projeto:
        produto_ids = request.POST.getlist('produto_ids')
        qtds = request.POST.getlist('qtds')
        count = 0
        for pid, qty in zip(produto_ids, qtds):
            try:
                produto = Produto.objects.get(pk=int(pid))
                ItemProjeto.objects.update_or_create(
                    projeto=projeto, produto=produto,
                    defaults={
                        'quantidade': int(qty) if qty else 1,
                        'preco_unitario': produto.unit_reais,
                    }
                )
                count += 1
            except (Produto.DoesNotExist, ValueError):
                pass

        if count > 0:
            projeto.revisao += 1
            projeto.save(update_fields=['revisao'])
        finalizar = request.POST.get('finalizar') == '1'
        if finalizar:
            return redirect('fluxo_projeto', projeto_id=projeto.id)
        return redirect('bom_selector_projeto', projeto_id=projeto.id)

    categoria = request.GET.get('categoria', '')
    busca = request.GET.get('busca', '')
    produtos = Produto.objects.all().order_by('categoria', 'nome')

    # Excluir produtos já incluídos neste projeto (exceto preco_variavel — sempre visível)
    if projeto:
        ids_ja_incluidos = (
            ItemProjeto.objects
            .filter(projeto=projeto, produto__preco_variavel=False)
            .values_list('produto_id', flat=True)
        )
        produtos = produtos.exclude(id_planilha__in=ids_ja_incluidos)

    if categoria:
        produtos = produtos.filter(categoria__iexact=categoria)
    if busca:
        produtos = produtos.filter(
            Q(nome__icontains=busca) | Q(modelo__icontains=busca) |
            Q(part_number__icontains=busca) | Q(fabricante__icontains=busca)
        )
    categorias = Produto.objects.values_list('categoria', flat=True).distinct().order_by('categoria')

    # Totais já incluídos no projeto
    total_itens_projeto = 0
    custo_projeto = 0
    if projeto:
        itens_projeto = ItemProjeto.objects.filter(projeto=projeto)
        total_itens_projeto = itens_projeto.count()
        custo_projeto = sum(i.preco_unitario * i.quantidade for i in itens_projeto)

    _mo_item = (
        ItemProjeto.objects.filter(projeto=projeto, produto__preco_variavel=True).first()
        if projeto else None
    )
    mo_preco = float(_mo_item.preco_unitario) if _mo_item else 0.0

    context = {
        'produtos': produtos, 'categorias': categorias,
        'categoria_ativa': categoria, 'busca': busca,
        'total_produtos': produtos.count(), 'projeto': projeto,
        'total_itens_projeto': total_itens_projeto,
        'custo_projeto': custo_projeto,
        'mo_preco': mo_preco,
    }
    return render(request, 'bom_selector.html', context)


# ============================================================
# FLUXO DO PROJETO (tela final do BoM construído)
# ============================================================
@login_required(login_url='/')
def fluxo_projeto(request, projeto_id):
    projeto = get_object_or_404(Projeto, id=projeto_id)

    if request.method == 'POST':
        acao = request.POST.get('acao', '')

        def _salvar_qtds():
            for item in ItemProjeto.objects.filter(projeto=projeto):
                nova_qty = request.POST.get(f'qty_{item.id}')
                if nova_qty and int(nova_qty) != item.quantidade:
                    item.quantidade = int(nova_qty)
                    item.save()

        def _incrementar_revisao():
            projeto.revisao += 1
            projeto.save(update_fields=['revisao'])

        # Excluir item
        if acao == 'excluir':
            item_id = request.POST.get('item_id')
            ItemProjeto.objects.filter(id=item_id, projeto=projeto).delete()
            _incrementar_revisao()
            return redirect('fluxo_projeto', projeto_id=projeto.id)

        # Salvar quantidades
        if acao == 'salvar_qtds':
            _salvar_qtds()
            _incrementar_revisao()
            return redirect('fluxo_projeto', projeto_id=projeto.id)

        # Salvar e sair (sem finalizar)
        if acao == 'salvar_sair':
            _salvar_qtds()
            _incrementar_revisao()
            return redirect('consultar_projeto')

        # Finalizar projeto
        if acao == 'finalizar':
            _salvar_qtds()
            _incrementar_revisao()
            projeto.finalizado = True
            projeto.save()
            return redirect('consultar_projeto')

    itens = ItemProjeto.objects.filter(projeto=projeto).select_related('produto')

    total_materiais = 0
    total_servico = 0
    for item in itens:
        valor = item.preco_unitario * item.quantidade
        if item.faturar_servico:
            total_servico += valor
        else:
            total_materiais += valor

    custo_total = total_materiais + total_servico
    venda_total = custo_total  # TODO: aplicar MKP

    context = {
        'projeto': projeto,
        'itens': itens,
        'total_materiais': total_materiais,
        'total_servico': total_servico,
        'custo_total': custo_total,
        'venda_total': venda_total,
    }
    return render(request, 'fluxo_projeto.html', context)


# ============================================================
# EXPORTAÇÃO DO PROJETO
# ============================================================

# Multiplicadores de locação: custo_total × fator = mensalidade
FATORES_LOCACAO = {24: 0.111111111, 36: 0.092165899, 48: 0.078740157, 60: 0.071684588}
FATOR_MANUTENCAO = 0.02 * 1.111   # manutenção = preço_venda × fator


def _numero_por_extenso(valor):
    """Converte valor Decimal/float para extenso em pt-BR."""
    try:
        from num2words import num2words
        reais = int(valor)
        centavos = round((float(valor) - reais) * 100)
        txt_reais = num2words(reais, lang='pt_BR').title()
        txt_real = 'Real' if reais == 1 else 'Reais'
        if centavos > 0:
            txt_cent = num2words(centavos, lang='pt_BR').title()
            txt_cent_label = 'Centavo' if centavos == 1 else 'Centavos'
            return f'{txt_reais} {txt_real} e {txt_cent} {txt_cent_label}'
        return f'{txt_reais} {txt_real}'
    except Exception:
        return str(valor)


def _formatar_brl(valor):
    from decimal import Decimal
    v = float(valor) if valor else 0.0
    return f'R$ {v:_.2f}'.replace('.', ',').replace('_', '.')


def _agrupar_infra(itens, desconto_pct=0):
    """Separa itens normais de infra; devolve (itens_normais, item_infra_ou_None)."""
    CATS_INFRA = {'INFRA', 'INFRAESTRUTURA', 'INFRA ESTRUTURA'}
    normais = []
    infra_total = 0
    for item in itens:
        cat = (item.produto.categoria or '').upper().strip()
        valor = float(item.preco_unitario) * item.quantidade
        if cat in CATS_INFRA:
            infra_total += valor
        else:
            normais.append(item)
    infra_item = None
    if infra_total > 0:
        desconto = desconto_pct / 100
        infra_item = {
            'nome': 'Infraestrutura',
            'modelo': '',
            'fabricante': '',
            'pn': '',
            'unidade': 'Vb',
            'quantidade': 1,
            'preco_unit': infra_total * (1 - desconto),
            'preco_total': infra_total * (1 - desconto),
        }
    return normais, infra_item


def _itens_para_linhas(itens, desconto_pct=0, incluir_infra=True):
    """Converte queryset em lista de dicts com desconto aplicado."""
    normais, infra = _agrupar_infra(itens, desconto_pct)
    linhas = []
    desconto = desconto_pct / 100
    seq = 1
    for item in normais:
        punit = float(item.preco_unitario) * (1 - desconto)
        ptotal = punit * item.quantidade
        linhas.append({
            'seq': seq,
            'id': item.produto.id_planilha,
            'nome': item.produto.nome,
            'modelo': item.produto.modelo or '',
            'fabricante': item.produto.fabricante or '',
            'pn': item.produto.part_number or '',
            'unidade': item.produto.unidade or 'Un',
            'quantidade': item.quantidade,
            'preco_unit': punit,
            'preco_total': ptotal,
            'servico': item.faturar_servico,
        })
        seq += 1
    if incluir_infra and infra:
        infra['seq'] = seq
        linhas.append(infra)
    return linhas


def _calcular_totais(itens, desconto_pct=0):
    desconto = desconto_pct / 100
    custo = sum(float(i.preco_unitario) * i.quantidade for i in itens)
    venda = custo * (1 - desconto)
    locacao = {p: custo * f for p, f in FATORES_LOCACAO.items()}
    manutencao = venda * FATOR_MANUTENCAO
    return {'custo': custo, 'venda': venda, 'desconto': desconto_pct,
            'locacao': locacao, 'manutencao': manutencao}


def _excel_tecnica(projeto, itens, desconto_pct, incluir_precos=False, modo='tecnica'):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Técnica' if not incluir_precos else 'Comercial'

    # Cores
    AZUL = '000666'
    BRANCO = 'FFFFFF'
    CINZA = 'F2F2F2'
    VERDE = 'E8F5E9'

    def celula(row, col, valor, bold=False, bg=None, cor=None, align='left', wrap=False, numero=False):
        c = ws.cell(row=row, column=col, value=valor)
        c.font = Font(bold=bold, color=cor or '000000', size=9)
        if bg:
            c.fill = PatternFill('solid', fgColor=bg)
        c.alignment = Alignment(horizontal=align, vertical='center', wrap_text=wrap)
        if numero and isinstance(valor, (int, float)):
            c.number_format = '#,##0.00'
        return c

    # Linha 1: Empresa
    empresa_nome = projeto.empresa_executora or 'PRISMA AXON'
    ws.merge_cells('A1:G1') if not incluir_precos else ws.merge_cells('A1:I1')
    c = ws.cell(row=1, column=1, value=f'{empresa_nome} — PROPOSTA TÉCNICA')
    c.font = Font(bold=True, color=BRANCO, size=12)
    c.fill = PatternFill('solid', fgColor=AZUL)
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 28

    # Linha 2: Info do projeto
    ws.merge_cells('A2:G2') if not incluir_precos else ws.merge_cells('A2:I2')
    c2 = ws.cell(row=2, column=1,
                 value=f'Projeto: {projeto.id_projeto_manual} | Cliente: {projeto.nome_cliente} | {projeto.municipio_obra}/{projeto.estado_obra} | {projeto.revisao_label}')
    c2.font = Font(bold=False, color='444444', size=9)
    c2.fill = PatternFill('solid', fgColor='E8EAF6')
    c2.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[2].height = 16

    # Linha 3: cabeçalho
    headers = ['Nº', 'Descrição', 'Modelo', 'Fabricante', 'P/N', 'Unid.', 'Qtd.']
    if incluir_precos:
        headers += ['Unit. (R$)', 'Total (R$)']
    for col, h in enumerate(headers, 1):
        celula(3, col, h, bold=True, bg=AZUL, cor=BRANCO, align='center')
    ws.row_dimensions[3].height = 18

    # Larguras
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 55
    ws.column_dimensions['C'].width = 22
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 7
    ws.column_dimensions['G'].width = 7
    if incluir_precos:
        ws.column_dimensions['H'].width = 14
        ws.column_dimensions['I'].width = 16

    linhas = _itens_para_linhas(itens, desconto_pct)
    for i, ln in enumerate(linhas):
        row = 4 + i
        bg = CINZA if i % 2 == 0 else BRANCO
        eh_infra = ln.get('nome') == 'Infraestrutura'
        if eh_infra:
            bg = VERDE
        celula(row, 1, ln['seq'], align='center', bg=bg)
        celula(row, 2, ln['nome'] + (f' — {ln["modelo"]}' if ln.get('modelo') else ''), bg=bg, wrap=True)
        celula(row, 3, ln.get('modelo', ''), bg=bg)
        celula(row, 4, ln.get('fabricante', ''), bg=bg)
        celula(row, 5, ln.get('pn', ''), bg=bg)
        celula(row, 6, ln.get('unidade', 'Un'), align='center', bg=bg)
        celula(row, 7, ln['quantidade'], align='center', bg=bg)
        if incluir_precos:
            celula(row, 8, ln['preco_unit'], align='right', bg=bg, numero=True)
            celula(row, 9, ln['preco_total'], align='right', bg=bg, numero=True)
        ws.row_dimensions[row].height = 28 if eh_infra else 15

    # Rodapé comercial
    if incluir_precos:
        totais = _calcular_totais(itens, desconto_pct)
        row_ini = 4 + len(linhas) + 1
        ws.merge_cells(f'A{row_ini}:H{row_ini}')
        ws.cell(row=row_ini, column=1, value='').fill = PatternFill('solid', fgColor=CINZA)

        def rod(r, label, valor, fmt_brl=True):
            ws.merge_cells(f'A{r}:H{r}')
            c = ws.cell(row=r, column=1, value=label)
            c.font = Font(bold=True, size=9, color='444444')
            c.alignment = Alignment(horizontal='right')
            cv = ws.cell(row=r, column=9, value=valor)
            cv.font = Font(bold=True, size=9, color=AZUL)
            cv.alignment = Alignment(horizontal='right')
            if fmt_brl and isinstance(valor, float):
                cv.number_format = '#,##0.00'

        rod(row_ini + 1, 'Custo Total:', totais['custo'])
        rod(row_ini + 2, f'Desconto ({desconto_pct}%):', totais['custo'] - totais['venda'])
        rod(row_ini + 3, 'Preço de Venda:', totais['venda'])
        r = row_ini + 4
        ws.merge_cells(f'A{r}:I{r}')
        ws.cell(row=r, column=1).fill = PatternFill('solid', fgColor=CINZA)

        rod2_start = r + 1
        ws.cell(row=rod2_start, column=1, value='Locação').font = Font(bold=True, size=9)
        for idx, (prazo, val) in enumerate(totais['locacao'].items()):
            rr = rod2_start + idx
            ws.merge_cells(f'A{rr}:H{rr}')
            ws.cell(row=rr, column=1, value=f'  {prazo} meses — mensalidade:').font = Font(size=9)
            cv = ws.cell(row=rr, column=9, value=val)
            cv.font = Font(bold=True, size=9, color='1B5E20')
            cv.number_format = '#,##0.00'
            cv.alignment = Alignment(horizontal='right')

        rr = rod2_start + len(totais['locacao'])
        ws.merge_cells(f'A{rr}:H{rr}')
        ws.cell(row=rr, column=1, value='  Manutenção (opcional) — mensalidade:').font = Font(size=9, italic=True)
        cv = ws.cell(row=rr, column=9, value=totais['manutencao'])
        cv.font = Font(bold=True, size=9, italic=True, color='E65100')
        cv.number_format = '#,##0.00'
        cv.alignment = Alignment(horizontal='right')

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _excel_projeto(projeto, itens):
    """Excel de template para importar projeto similar."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    import io

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Projeto'

    AZUL = '000666'
    BRANCO = 'FFFFFF'
    CINZA = 'F2F2F2'

    headers = ['ID Produto', 'Nome', 'Modelo', 'P/N', 'Fabricante', 'Unidade', 'Quantidade', 'Preço Unit (R$)', 'Preço Total (R$)']
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = Font(bold=True, color=BRANCO, size=9)
        c.fill = PatternFill('solid', fgColor=AZUL)
        c.alignment = Alignment(horizontal='center', vertical='center')

    larguras = [12, 50, 22, 16, 16, 8, 8, 16, 16]
    for i, w in enumerate(larguras, 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w

    for row, item in enumerate(itens, 2):
        bg = CINZA if row % 2 == 0 else BRANCO
        vals = [
            item.produto.id_planilha,
            item.produto.nome,
            item.produto.modelo or '',
            item.produto.part_number or '',
            item.produto.fabricante or '',
            item.produto.unidade or 'Un',
            item.quantidade,
            float(item.preco_unitario),
            float(item.preco_unitario) * item.quantidade,
        ]
        for col, val in enumerate(vals, 1):
            c = ws.cell(row=row, column=col, value=val)
            c.font = Font(size=9)
            c.fill = PatternFill('solid', fgColor=bg)
            if col in (8, 9) and isinstance(val, float):
                c.number_format = '#,##0.00'
                c.alignment = Alignment(horizontal='right')

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _texto_extenso(projeto, itens, desconto_pct):
    totais = _calcular_totais(itens, desconto_pct)
    cliente = projeto.nome_cliente.upper()
    estado = projeto.estado_obra.upper()
    empresa = projeto.empresa_executora or 'EMIVE'

    v_capex = totais['venda']
    v_manut = totais['manutencao']
    v_36 = totais['locacao'][36]
    v_60 = totais['locacao'][60]

    def brl(v): return _formatar_brl(v)
    def ext(v): return _numero_por_extenso(round(v, 2))

    return f"""11.\tValores da Nossa Proposta

Para atender de forma ideal às necessidades e à estratégia financeira da {cliente}, apresentamos a seguir duas modalidades de contratação:

Aquisição (CAPEX): Esta opção permite a compra completa dos equipamentos e da infraestrutura. É um modelo de investimento de capital que oferece a propriedade total do sistema. A proposta de aquisição pode ser complementada com um contrato de manutenção opcional, garantindo o suporte técnico e a preservação do sistema após a implantação.

Locação (OPEX): Um modelo de serviço completo, com contratos de 36 ou 60 meses, que inclui a instalação, os equipamentos, a manutenção e o suporte técnico e inclusive troca de equipamentos em caso de defeito ao longo do contrato. Essa opção permite que a {cliente} transforme um custo de capital (CAPEX) em uma despesa operacional (OPEX), liberando recursos para outros investimentos e garantindo que o sistema esteja sempre atualizado e em pleno funcionamento.

Esses modelos foram elaborados para oferecer a máxima flexibilidade, permitindo à {cliente} selecionar a opção que melhor se alinhe com seus objetivos estratégicos e planejamento orçamentário.

Aquisição do Sistema (CAPEX).
I.\tValor para a Aquisição de Todos os Equipamentos e Sistema de Software Licenciado (inclui garantia balcão de 12 meses dos equipamentos instalados em caso de defeitos de fábrica):
{brl(v_capex)} ({ext(v_capex)})
*Condições Comerciais Item 13.

II.\tOPCIONAL » Valor Mensal para Manutenção Preventiva e Corretiva (Inclui substituição de peças e acessórios, limpeza, atualizações de Hardware e Software, Suporte Técnico, Manutenção Preventiva, Manutenção Corretiva, deslocamentos e visitas técnicas de acordo com SLA):
{brl(v_manut)} ({ext(v_manut)})
*Este Valor deverá ser disponibilizado mensalmente pela Manutenção dos Equipamentos.


Locação do Sistema (OPEX).

36 MESES
I.\tValor Mensal para Locação dos Novos Equipamentos de todas as áreas do CD {cliente}/{estado} (Inclui Garantia Total do Sistema, Substituição de Peças, Atualizações de Hardware e Software, Suporte Técnico, Manutenção Preventiva, Manutenção Corretiva):
{brl(v_36)} ({ext(v_36)}).
*Condições Comerciais no item 14.

60 MESES
II.\tValor Mensal para Locação dos Novos Equipamentos de todas as áreas do CD {cliente}/{estado} (Inclui Garantia Total do Sistema, Substituição de Peças, Atualizações de Hardware e Software, Suporte Técnico, Manutenção Preventiva, Manutenção Corretiva):
{brl(v_60)} ({ext(v_60)}).
*Condições Comerciais no item 14.

Benefícios da solução LOCAÇÃO {empresa}:
\t• Equipamentos sempre atualizados;
\t• Economia com suporte técnico;
\t• Garantia total de todos os equipamentos cobertos pelo contrato;
\t• Não há depreciação, desvalorização ou obsolescência dos equipamentos;
\t• Fim do problema com a vida útil dos equipamentos (Equipamentos eletrônicos de 3 a 5 anos);
\t• Benefícios fiscais para empresas no regime de lucro real (Abatimento no IRPJ e CSLL);
\t• Previsibilidade do investimento para manutenção de todo o parque de equipamentos;
\t• Equipamentos de backup para agilidade das manutenções e reparos;
\t• Soluções dos problemas e defeitos dos equipamentos em um tempo muito mais rápido;
"""


@login_required(login_url='/')
def exportar_projeto(request, projeto_id):
    """Endpoint único de exportação. Parâmetros: modo (tecnica/comercial/locacao/extenso/projeto/filtro) e fmt (xlsx/docx)."""
    from django.http import HttpResponse
    import io

    projeto = get_object_or_404(Projeto, pk=projeto_id)
    modo = request.GET.get('modo', 'tecnica')
    fmt = request.GET.get('fmt', 'xlsx')
    desconto_pct = float(request.GET.get('desconto', '0') or '0')

    itens = list(ItemProjeto.objects.filter(projeto=projeto).select_related('produto').order_by('produto__categoria', 'produto__nome'))
    nome_arquivo = f'{projeto.id_projeto_manual}_{modo}'

    # ── Excel ──────────────────────────────────────────────────────────────
    if fmt == 'xlsx':
        if modo in ('tecnica', 'filtro'):
            buf = _excel_tecnica(projeto, itens, desconto_pct, incluir_precos=False)
        elif modo == 'comercial':
            buf = _excel_tecnica(projeto, itens, desconto_pct, incluir_precos=True)
        elif modo == 'locacao':
            buf = _excel_tecnica(projeto, itens, desconto_pct, incluir_precos=True)
        elif modo == 'projeto':
            buf = _excel_projeto(projeto, itens)
        elif modo == 'extenso':
            # Extenso em Excel = planilha de locação/comercial sem o texto
            buf = _excel_tecnica(projeto, itens, desconto_pct, incluir_precos=True)
        else:
            buf = _excel_tecnica(projeto, itens, desconto_pct, incluir_precos=False)

        resp = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        resp['Content-Disposition'] = f'attachment; filename="{nome_arquivo}.xlsx"'
        return resp

    # ── Word ───────────────────────────────────────────────────────────────
    if fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        # Margens menores
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Cabeçalho
        h = doc.add_heading(f'{projeto.empresa_executora or "PRISMA AXON"} — PROPOSTA TÉCNICA', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info = doc.add_paragraph(f'Projeto: {projeto.id_projeto_manual} | Cliente: {projeto.nome_cliente} | {projeto.municipio_obra}/{projeto.estado_obra} | {projeto.revisao_label}')
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if modo == 'extenso':
            texto = _texto_extenso(projeto, itens, desconto_pct)
            for linha in texto.split('\n'):
                p = doc.add_paragraph(linha)
                p.paragraph_format.space_after = Pt(2)
        else:
            incluir_precos = modo in ('comercial', 'locacao')
            linhas = _itens_para_linhas(itens, desconto_pct)
            headers = ['Nº', 'Descrição', 'Modelo', 'Fabricante', 'Qtd.']
            if incluir_precos:
                headers += ['Unit. (R$)', 'Total (R$)']

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h_txt in enumerate(headers):
                hdr_cells[i].text = h_txt
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

            for ln in linhas:
                row_cells = table.add_row().cells
                vals = [str(ln['seq']), ln['nome'], ln.get('modelo',''), ln.get('fabricante',''), str(ln['quantidade'])]
                if incluir_precos:
                    vals += [_formatar_brl(ln['preco_unit']).replace('R$ ',''), _formatar_brl(ln['preco_total']).replace('R$ ','')]
                for i, v in enumerate(vals):
                    row_cells[i].text = v
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

            if incluir_precos:
                totais = _calcular_totais(itens, desconto_pct)
                doc.add_paragraph()
                for label, val in [
                    ('Custo Total', totais['custo']),
                    (f'Desconto ({desconto_pct}%)', totais['custo'] - totais['venda']),
                    ('Preço de Venda', totais['venda']),
                ]:
                    p = doc.add_paragraph()
                    p.add_run(f'{label}: ').bold = True
                    p.add_run(_formatar_brl(val))

                doc.add_paragraph('Locação Mensal:').runs[0].bold = True
                for prazo, val in totais['locacao'].items():
                    doc.add_paragraph(f'  {prazo} meses: {_formatar_brl(val)}')
                doc.add_paragraph(f'  Manutenção (opcional): {_formatar_brl(totais["manutencao"])}')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = f'attachment; filename="{nome_arquivo}.docx"'
        return resp

    return HttpResponse('Formato inválido', status=400)


# ============================================================
# CANCELAR PRÉ-PROJETO
# ============================================================
@login_required(login_url='/')
def cancelar_projeto(request, projeto_id):
    projeto = get_object_or_404(Projeto, id=projeto_id)
    if not projeto.finalizado:
        projeto.delete()
    return redirect('inicio')


# ============================================================
# CADASTRO DE PRODUTO
# ============================================================
@login_required(login_url='/')
def cadastro_produto(request):
    if request.method == 'POST':
        categoria = request.POST.get('categoria', 'OUTROS')
        novo_id = gerar_proximo_id(categoria)
        try:
            Produto.objects.create(
                id_planilha=novo_id, nome=request.POST.get('nome', ''),
                modelo=request.POST.get('modelo', ''), part_number=request.POST.get('part_number', ''),
                fabricante=request.POST.get('fabricante', ''), unidade=request.POST.get('unidade', 'peça'),
                categoria=categoria, descricao=request.POST.get('descricao', ''),
                moeda=request.POST.get('moeda', 'reais'),
                preco_fornecedor=request.POST.get('preco_fornecedor', 0) or 0,
                unit_reais=request.POST.get('unit_reais', 0) or 0,
                frete_na_compra=request.POST.get('frete_na_compra', 0) or 0,
                ipi=request.POST.get('ipi', 0) or 0, icms=request.POST.get('icms', 0) or 0,
                nome_fornecedor=request.POST.get('nome_fornecedor', ''),
                estado_origem=request.POST.get('estado_origem', ''),
                grupo_financeiro=request.POST.get('grupo_financeiro', ''),
                ncm=request.POST.get('ncm', ''),
                data_ultima_cotacao=request.POST.get('data_ultima_cotacao') or None,
                lucro_percent=request.POST.get('lucro_percent', 0) or 0,
                iss_percent=request.POST.get('iss_percent', 0) or 0,
                pis_cofins_percent=request.POST.get('pis_cofins_percent', 0) or 0,
                ir_csll_lp=request.POST.get('ir_csll_lp', 0) or 0,
                ir_csll_lr=request.POST.get('ir_csll_lr', 0) or 0,
                mkp=request.POST.get('mkp', 0) or 0,
                custo_loc=request.POST.get('custo_loc', 0) or 0,
                custo_mensal=request.POST.get('custo_mensal', 0) or 0,
                iss_loc=request.POST.get('iss_loc', 0) or 0,
                pis_cofins_loc=request.POST.get('pis_cofins_loc', 0) or 0,
                ir_csll_lp_loc=request.POST.get('ir_csll_lp_loc', 0) or 0,
                ir_csll_lr_loc=request.POST.get('ir_csll_lr_loc', 0) or 0,
                mkp_loc=request.POST.get('mkp_loc', 0) or 0,
                status='amarelo',  # Novo produto = aguardando validação
            )
            return redirect('inicio')
        except Exception as e:
            messages.error(request, f'Erro: {e}')
    return render(request, 'cadastro_produto.html', {'modo': 'criar'})


@login_required(login_url='/')
def detalhe_produto(request, pk):
    produto = get_object_or_404(Produto, pk=pk)
    if request.method == 'POST':
        for field in ['nome','modelo','part_number','fabricante','unidade','categoria','descricao','moeda',
                       'nome_fornecedor','estado_origem','grupo_financeiro','ncm']:
            setattr(produto, field, request.POST.get(field, getattr(produto, field)))
        for field in ['preco_fornecedor','unit_reais','frete_na_compra','ipi','icms',
                       'lucro_percent','iss_percent','pis_cofins_percent','ir_csll_lp','ir_csll_lr','mkp',
                       'custo_loc','custo_mensal','iss_loc','pis_cofins_loc','ir_csll_lp_loc','ir_csll_lr_loc','mkp_loc']:
            setattr(produto, field, request.POST.get(field, getattr(produto, field)) or 0)
        try:
            produto.save()
            return redirect('detalhe_produto', pk=produto.pk)
        except Exception as e:
            messages.error(request, f'Erro: {e}')
    return render(request, 'cadastro_produto.html', {'produto': produto, 'modo': 'detalhe'})


# ============================================================
# GESTÃO DE USUÁRIOS
# ============================================================
@login_required(login_url='/')
def gestao_usuarios(request):
    usuarios = User.objects.all().order_by('username')
    return render(request, 'gestao_usuarios.html', {'usuarios': usuarios})


# ============================================================
# GESTÃO DE EMPRESAS
# ============================================================
@login_required(login_url='/')
def gestao_empresas(request):
    empresas = Empresa.objects.all().order_by('nome_fantasia')
    return render(request, 'gestao_empresas.html', {'empresas': empresas})


@login_required(login_url='/')
def cadastro_empresa(request):
    if request.method == 'POST':
        cnpj = request.POST.get('cnpj', '').strip()
        if Empresa.objects.filter(cnpj=cnpj).exists():
            messages.error(request, f'Já existe uma empresa com o CNPJ "{cnpj}".')
            return render(request, 'cadastro_empresa.html')
        try:
            empresa = Empresa(
                razao_social=request.POST.get('razao_social', ''),
                nome_fantasia=request.POST.get('nome_fantasia', ''),
                cnpj=cnpj,
            )
            if request.FILES.get('logo'):
                empresa.logo = request.FILES['logo']
            empresa.save()
            return redirect('gestao_empresas')
        except Exception as e:
            messages.error(request, f'Erro ao cadastrar empresa: {e}')
    return render(request, 'cadastro_empresa.html')


@login_required(login_url='/')
def editar_empresa(request, empresa_id):
    edit_empresa = get_object_or_404(Empresa, id=empresa_id)
    if request.method == 'POST':
        edit_empresa.razao_social = request.POST.get('razao_social', edit_empresa.razao_social)
        edit_empresa.nome_fantasia = request.POST.get('nome_fantasia', edit_empresa.nome_fantasia)
        if request.FILES.get('logo'):
            edit_empresa.logo = request.FILES['logo']
        try:
            edit_empresa.save()
            return redirect('gestao_empresas')
        except Exception as e:
            messages.error(request, f'Erro: {e}')
    return render(request, 'cadastro_empresa.html', {'edit_empresa': edit_empresa})


# ============================================================
# CADASTRO / EDIÇÃO DE USUÁRIO
# ============================================================
@login_required(login_url='/')
def cadastro_usuario(request):
    empresas = Empresa.objects.all().order_by('nome_fantasia')
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        password = request.POST.get('password', '')
        password2 = request.POST.get('password2', '')

        if not username:
            messages.error(request, 'Nome de usuário é obrigatório.')
            return render(request, 'cadastro_usuario.html', {'empresas': empresas})
        if User.objects.filter(username=username).exists():
            messages.error(request, f'O nome de usuário "{username}" já existe.')
            return render(request, 'cadastro_usuario.html', {'empresas': empresas})
        if password != password2:
            messages.error(request, 'As senhas não conferem.')
            return render(request, 'cadastro_usuario.html', {'empresas': empresas})
        if len(password) < 4:
            messages.error(request, 'A senha deve ter pelo menos 4 caracteres.')
            return render(request, 'cadastro_usuario.html', {'empresas': empresas})

        try:
            nivel = request.POST.get('nivel', 'analista')
            user = User.objects.create_user(
                username=username, email=email, password=password,
                first_name=first_name, last_name=last_name,
            )
            if nivel == 'admin':
                user.is_superuser = True
                user.is_staff = True
            elif nivel in ('gerente', 'supervisor'):
                user.is_staff = True
            user.save()

            # Criar perfil com foto, cargo e empresa
            empresa_id = request.POST.get('empresa', '')
            perfil = Perfil.objects.create(
                user=user,
                cargo=request.POST.get('cargo', 'analista'),
                empresa_id=int(empresa_id) if empresa_id else None,
            )
            if request.FILES.get('foto'):
                perfil.foto = request.FILES['foto']
                perfil.save()

            return redirect('gestao_usuarios')
        except Exception as e:
            messages.error(request, f'Erro ao criar usuário: {e}')

    return render(request, 'cadastro_usuario.html', {'empresas': empresas})


@login_required(login_url='/')
def editar_usuario(request, user_id):
    edit_user = get_object_or_404(User, id=user_id)
    empresas = Empresa.objects.all().order_by('nome_fantasia')

    # Garantir que o perfil existe
    perfil, _ = Perfil.objects.get_or_create(user=edit_user)

    if request.method == 'POST':
        edit_user.first_name = request.POST.get('first_name', edit_user.first_name)
        edit_user.last_name = request.POST.get('last_name', edit_user.last_name)
        edit_user.email = request.POST.get('email', edit_user.email)

        nivel = request.POST.get('nivel', 'analista')
        edit_user.is_superuser = (nivel == 'admin')
        edit_user.is_staff = (nivel in ('admin', 'supervisor'))

        perfil.cargo = request.POST.get('cargo', perfil.cargo)
        empresa_id = request.POST.get('empresa', '')
        perfil.empresa_id = int(empresa_id) if empresa_id else None

        if request.FILES.get('foto'):
            try:
                from google.cloud import storage as _gcs
                import re, uuid, os
                _arq = request.FILES['foto']
                _ext = os.path.splitext(_arq.name)[1].lower() or '.jpg'
                _safe = re.sub(r'[^a-zA-Z0-9]', '_', os.path.splitext(_arq.name)[0])[:40]
                _blob_name = f'media/usuarios/{_safe}_{uuid.uuid4().hex[:8]}{_ext}'
                _bucket_name = 'primaaxon-media-files'
                _client = _gcs.Client()
                _bucket = _client.bucket(_bucket_name)
                _blob = _bucket.blob(_blob_name)
                _blob.upload_from_file(_arq, content_type=_arq.content_type)
                print(f'[GCS DIRECT] Salvo em gs://{_bucket_name}/{_blob_name}', flush=True)
                perfil.foto = _blob_name.replace('media/', '', 1)  # remove prefixo media/ para o campo
                perfil.save(update_fields=['foto'])
                messages.success(request, f'Foto salva!')
            except Exception as e:
                import traceback
                print(f'[GCS DIRECT] ERRO: {e}\n{traceback.format_exc()}', flush=True)
                messages.error(request, f'Erro ao salvar foto: {e}')

        try:
            edit_user.save()
            perfil.save()
            return redirect('gestao_usuarios')
        except Exception as e:
            messages.error(request, f'Erro: {e}')

    return render(request, 'cadastro_usuario.html', {'edit_user': edit_user, 'perfil': perfil, 'empresas': empresas})


# ============================================================
# PRODUTO — ABAS DETALHADAS
# ============================================================
@login_required(login_url='/')
def produto_aba(request, pk, aba):
    produto = get_object_or_404(Produto, pk=pk)

    CAMPOS_TEXTO = {
        'cadastro': ['nome', 'modelo', 'part_number', 'fabricante', 'unidade', 'categoria', 'descricao'],
        'fiscal':   ['moeda'],
        'compras':  ['nome_fornecedor', 'estado_origem', 'grupo_financeiro', 'ncm'],
        'mkp':      [],
        'locacao':  [],
    }
    CAMPOS_DEC = {
        'cadastro': [],
        'fiscal':   ['preco_fornecedor', 'unit_reais', 'frete_na_compra', 'ipi', 'icms'],
        'compras':  [],
        'mkp':      ['lucro_percent', 'iss_percent', 'pis_cofins_percent', 'ir_csll_lp', 'ir_csll_lr', 'mkp'],
        'locacao':  ['custo_loc', 'custo_mensal', 'iss_loc', 'pis_cofins_loc', 'ir_csll_lp_loc', 'ir_csll_lr_loc', 'mkp_loc'],
    }
    TEMPLATES = {
        'cadastro': 'cadastro_produtos_inclusao.html',
        'fiscal':   'cadastro_produtos_fiscal.html',
        'compras':  'cadastro_produtos_compras.html',
        'mkp':      'cadastro_produtos_mkp.html',
        'locacao':  'cadastro_produtos_mkp_locacao.html',
    }

    if aba not in TEMPLATES:
        from django.http import Http404
        raise Http404

    if request.method == 'POST':
        for field in CAMPOS_TEXTO.get(aba, []):
            val = request.POST.get(field, '').strip()
            setattr(produto, field, val)
        for field in CAMPOS_DEC.get(aba, []):
            val = request.POST.get(field)
            setattr(produto, field, val or 0)
        if aba == 'compras':
            dt = request.POST.get('data_ultima_cotacao')
            if dt:
                produto.data_ultima_cotacao = dt
        if aba == 'cadastro':
            produto.preco_variavel = request.POST.get('preco_variavel') == 'on'
            if request.FILES.get('foto'):
                produto.foto = request.FILES['foto']
        try:
            produto.save()
            messages.success(request, 'Dados salvos.')
        except Exception as e:
            messages.error(request, f'Erro: {e}')
        return redirect('produto_aba', pk=pk, aba=aba)

    return render(request, TEMPLATES[aba], {'produto': produto, 'aba': aba})


# ============================================================
# GESTÃO DO PROGRAMA (só Prisma Axon sistema)
# ============================================================
@login_required(login_url='/')
def gestao_programa(request):
    from django.http import JsonResponse
    # Verificar se é usuário do sistema
    try:
        is_sistema = request.user.perfil.empresa.is_sistema
    except Exception:
        is_sistema = request.user.is_superuser
    if not is_sistema:
        messages.error(request, 'Acesso restrito.')
        return redirect('inicio')

    # Verificar senha via POST
    if request.method == 'POST' and request.POST.get('acao') == 'verificar_senha':
        senha = request.POST.get('senha', '')
        user = authenticate(request, username=request.user.username, password=senha)
        if user is not None:
            return JsonResponse({'ok': True})
        return JsonResponse({'ok': False})

    return render(request, 'gestão_do_programa.html', {'is_sistema': is_sistema})


# ============================================================
# RESET DE SENHA DE USUÁRIO
# ============================================================
@login_required(login_url='/')
def reset_senha_usuario(request, user_id):
    from django.utils import timezone
    edit_user = get_object_or_404(User, id=user_id)
    try:
        empresa = edit_user.perfil.empresa
        nome_empresa = empresa.nome_fantasia.split()[0] if empresa else 'Prisma'
    except Exception:
        nome_empresa = 'Prisma'
    ano = timezone.now().year
    nova_senha = f"{nome_empresa}@{ano}"
    edit_user.set_password(nova_senha)
    edit_user.save()
    messages.success(request, f'Senha redefinida para: {nova_senha}')
    return redirect('editar_usuario', user_id=user_id)


# ============================================================
# TROCAR SENHA PRÓPRIA (AJAX)
# ============================================================
@login_required(login_url='/')
def trocar_senha_proprio(request):
    from django.http import JsonResponse
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'erro': 'Método inválido'})
    senha_atual = request.POST.get('senha_atual', '')
    senha_nova = request.POST.get('senha_nova', '')
    user = authenticate(request, username=request.user.username, password=senha_atual)
    if user is None:
        return JsonResponse({'ok': False, 'erro': 'Senha atual incorreta.'})
    if len(senha_nova) < 4:
        return JsonResponse({'ok': False, 'erro': 'A nova senha deve ter ao menos 4 caracteres.'})
    user.set_password(senha_nova)
    user.save()
    # Mantém a sessão ativa após a troca
    from django.contrib.auth import update_session_auth_hash
    update_session_auth_hash(request, user)
    return JsonResponse({'ok': True})


# ============================================================
# ENTRAR NO M.O — auto-inclui preco_variavel no BoM e redireciona
# ============================================================
@login_required(login_url='/')
def entrar_mo(request, projeto_id):
    from decimal import Decimal
    projeto = get_object_or_404(Projeto, pk=projeto_id)
    # Garante que todos os produtos preco_variavel estão no BoM (remarca se foi desmarcado)
    for produto in Produto.objects.filter(preco_variavel=True):
        item, criado = ItemProjeto.objects.get_or_create(
            projeto=projeto,
            produto=produto,
            defaults={'quantidade': 1, 'preco_unitario': Decimal('0.00')},
        )
        # Se já existia mas foi removido e recriado, não mexe no preco_unitario existente
    return redirect('gestao_mo', projeto_id=projeto_id, aba='operacional')


# ============================================================
# GESTÃO DE MÃO DE OBRA (abas por projeto)
# ============================================================
ABAS_MO = [
    ('operacional',   'Mão de Obra',          'engineering'),
    ('ferramentas',   'Equipamentos/Insumos',  'construction'),
    ('transporte',    'Transporte',             'local_shipping'),
    ('demais_custos', 'Demais Custos',          'payments'),
    ('terceiros',     'Terceirizados',          'handshake'),
]

# Encargos sociais aplicados sobre o custo de M.O (INSS patronal + FGTS + outros)
ENCARGOS_PERCENT = 77.92  # %

# Salários base mensais (S/ADC — sem adicionais) usados para auto-fill de custo
# Dias = mensal / 20  |  Horas = mensal / 146
SALARIOS_MO = {
    'Auxiliar de Instalação':              1780.16,
    'Dupla Técnica':                       4582.56,
    'Dupla Técnica (Civil)':               4582.56,
    'Eletricista':                         1827.50,
    'Técnico de Manutenção':               2802.39,
    'Técnico de TI':                       3654.40,
    'Técnico de T.I':                      3654.40,
    'Técnico de TI - Ferroport':           2922.51,
    'Técnico de Manutenção - Ferroport':   2734.67,
    'Técnico de Manutenção - Prumo':       1979.58,
    'Técnico de Segurança do Trabalho':    3008.00,
    'Coordenador Técnico':                 3993.95,
    'Técnico de Manutenção - Samarco':     2719.68,
    'Supervisor Técnico':                  5000.00,
    'Supervisor':                          5000.00,
    'Projetista':                          4000.00,
    'Engenheiro Eletricista':              7500.00,
    'Engenheiro Residente':                7500.00,
    'Pedreiro':                            2000.00,
    'Ajudante de Pedreiro':                1485.00,
}

FUNCAO_ICONS = {
    'Engenheiro Residente': 'engineering',
    'Engenheiro Eletricista': 'engineering',
    'Coordenador Técnico': 'manage_accounts',
    'Supervisor': 'supervisor_account',
    'Supervisor Técnico': 'supervisor_account',
    'Projetista': 'architecture',
    'Técnico de T.I': 'computer',
    'Técnico de TI': 'computer',
    'Eletricista': 'electrical_services',
    'Dupla Técnica': 'groups',
    'Dupla Técnica (Civil)': 'groups',
    'Técnico de Manutenção': 'build',
    'Pedreiro': 'construction',
    'Auxiliar de Instalação': 'handyman',
}

FROTA_TRANSPORTE = [
    ('Veículo Passeio',  'directions_car',   2700),
    ('Utilitário',       'airport_shuttle',   3500),
    ('Caçamba',          'local_shipping',    4000),
    ('Veículo Próprio',  'local_taxi',        1500),
    ('Motocicleta',      'two_wheeler',        800),
]

ITENS_DEMAIS_CUSTOS_DEFAULT = [
    # (descricao, _unused, custo_unit, unidade, tempo_default, ativo_default)
    # ativo_default=False → só ativa se houver viagem longa no Transporte
    ('Plano de Saúde',                   'Meses', 175.00, 'Meses', 12, True),
    ('Plano Odontológico',               'Meses',  12.00, 'Meses', 12, True),
    ('Seguro de Vida',                   'Meses',  12.00, 'Meses', 12, True),
    ('Uniforme',                         'Meses',  31.67, 'Meses',  1, True),
    ('Uniforme NR10',                    'Meses',  63.33, 'Meses',  1, True),
    ('Alimentação Convencional Técnico', 'Meses', 440.00, 'Meses', 12, True),
    ('Alimentação (Diária)',             'Dias',   20.00, 'Dias',  22, False),
    ('Hospedagem',                       'Dias',  220.00, 'Dias',  15, False),
    ('Pedágio',                          'Dias',    0.00, 'Dias',  22, False),
    ('PPRA/PCMSO',                       'Meses',  45.00, 'Meses',  1, True),
]

ITENS_TERCEIROS_DEFAULT = [
    # (descricao, especificacao, custo_unit)
    ('Instalação de Ponto de Câmera',                   'Por ponto',     80.00),
    ('Instalação de Ponto de Alarme',                   'Por sensor',    80.00),
    ('Instalação de Ponto de Controle de Acesso',       'Por leitor',    80.00),
    ('Instalação de Base de Cancela / Catraca',         'Por unidade',  400.00),
    ('Instalação de Poste Plantado',                    'Poste 6m',     500.00),
    ('Instalação de Poste com Base',                    'Poste+base',   700.00),
    ('Instalação e Montagem de Painel',                 'Por painel',   160.00),
    ('Instalação de Infra Aparente (m)',                'Por metro',     30.00),
    ('Instalação de Infra Subterrânea (m)',             'Por metro',     51.00),
    ('Instalação de Infra Subterrânea Envelopada (m)',  'Por metro',    120.00),
    ('Lançamento de Cabeamento em Tubulação (m)',       'Por metro',      1.50),
    ('Lançamento de Cabeamento Aéreo (m)',              'Por metro',      3.00),
    ('Alimentação (Diária)',                            'Por dia',       75.00),
    ('Hospedagem',                                      'Por dia',      220.00),
    ('Km de Deslocamento',                              'Por km',         1.40),
    ('Deslocamento Aéreo',                              'Passagem',       0.00),
    ('Deslocamento Ônibus',                             'Passagem',       0.00),
]

FUNCOES_MO = {
    'operacional': [
        'Auxiliar de Instalação', 'Dupla Técnica', 'Dupla Técnica (Civil)',
        'Eletricista', 'Técnico de Manutenção', 'Técnico de TI',
        'Técnico de TI - Ferroport', 'Técnico de Manutenção - Ferroport',
        'Técnico de Manutenção - Prumo', 'Técnico de Segurança do Trabalho',
        'Coordenador Técnico', 'Técnico de Manutenção - Samarco',
        'Supervisor Técnico', 'Projetista', 'Engenheiro Eletricista',
        'Engenheiro Residente', 'Pedreiro', 'Ajudante de Pedreiro',
    ],
    'ferramentas':   ['Maleta de Ferramentas', 'Equipamento de Medição', 'Insumos Perecíveis', 'EPI', 'Andaime/Escada', 'Escritório de Obra'],
    'transporte':    ['Veículo Leve', 'Van/Utilitário', 'Caminhão', 'Combustível', 'Pedágio', 'Hospedagem', 'Passagem Aérea'],
    'demais_custos': ['Vale Refeição', 'Vale Transporte', 'Plano de Saúde', 'Plano Odontológico', 'Seguro de Vida', 'Uniforme', 'Treinamento/Capacitação'],
    'terceiros':     ['Instalação Civil', 'Cabeamento Estruturado', 'Montagem de Racks', 'Comissionamento', 'Certificação de Rede', 'Projeto Executivo'],
}

CONFIG_ABA = {
    'operacional': {
        'col_descricao': 'Função / Cargo',
        'col_tempo': 'Tempo',
        'col_custo': 'Salário (R$/mês)',
        'unidade_default': 'Meses',
        'placeholder': 'Ex: Especialista em CFTV',
        'add_label': 'Adicionar Profissional',
        'descricao_label': 'Função / Cargo',
        'dica': 'Profissionais alocados ao projeto. Qtd = número de pessoas, Tempo = duração em meses.',
    },
    'ferramentas': {
        'col_descricao': 'Equipamento / Item',
        'col_tempo': 'Período',
        'col_custo': 'Custo Unit. (R$)',
        'unidade_default': 'Meses',
        'placeholder': 'Ex: Serra mármore, Multímetro, EPI...',
        'add_label': 'Adicionar Equipamento',
        'descricao_label': 'Equipamento / Item',
        'dica': 'Equipamentos, ferramentas e insumos consumidos. Qtd = quantidade, Período = meses de uso.',
    },
    'transporte': {
        'col_descricao': 'Tipo de Transporte / Item',
        'col_tempo': 'Qtd Viagens/Períodos',
        'col_custo': 'Custo Unit. (R$)',
        'unidade_default': 'Dias',
        'placeholder': 'Ex: Passagem Manaus–Brasília...',
        'add_label': 'Adicionar Transporte',
        'descricao_label': 'Tipo de Transporte',
        'dica': 'Veículos, passagens, hospedagem e pedágios. Qtd = nº de pessoas/veículos, Tempo = dias ou viagens.',
    },
    'demais_custos': {
        'col_descricao': 'Benefício / Item',
        'col_tempo': 'Período (Meses)',
        'col_custo': 'Custo Unit. (R$/mês)',
        'unidade_default': 'Meses',
        'placeholder': 'Ex: Cesta básica, Seguro específico...',
        'add_label': 'Adicionar Item',
        'descricao_label': 'Benefício / Item',
        'dica': 'Benefícios e encargos adicionais. Qtd = nº de beneficiários, Tempo = meses de vigência.',
    },
    'terceiros': {
        'col_descricao': 'Serviço / Fornecedor',
        'col_tempo': 'Etapas / Parcelas',
        'col_custo': 'Valor Unitário (R$)',
        'unidade_default': 'Vb',
        'placeholder': 'Ex: Empresa de cabeamento estruturado...',
        'add_label': 'Adicionar Serviço',
        'descricao_label': 'Serviço / Fornecedor',
        'dica': 'Serviços subcontratados. Qtd = nº de contratos/unidades, Tempo = etapas ou Vb (verba).',
    },
}

# ── Kits de Ferramentas ──────────────────────────────────────────────────────
# (nome, grupo, custo_mensal_unit)
KIT_MALETA_ITEMS = [
    ('Detector Gas Multigas Lel O2 Co H2s-Lc', 'Equipamento', 209.50),
    ('Crowcon T4 Portable 4-Gas Detector', 'Equipamento', 136.89),
    ('Monitor portátil LCD 3.5"', 'Equipamento', 6.94),
    ('Notebook', 'Equipamento', 69.44),
    ('Celular', 'Equipamento', 55.56),
    ('Kit localizador e testador de cabo', 'Equipamento', 4.17),
    ('Maquita', 'Equipamento', 11.11),
    ('Multímetro digital cat III', 'Ferramental', 4.44),
    ('Alicate Amperímetro', 'Ferramental', 6.83),
    ('Parafusadeira industrial elétrica ou a bateria', 'Ferramental', 31.73),
    ('Philips e tork', 'Ferramental', 1.21),
    ('Trena digital', 'Ferramental', 13.06),
    ('Furadeira', 'Ferramental', 11.11),
    ('Alicate de bico 6"', 'Ferramental', 0.98),
    ('Alicate de corte diagonal 6"', 'Ferramental', 0.98),
    ('Alicate de crimpar terminal tubular 0,5A 6mm Yac', 'Ferramental', 3.37),
    ('Alicate de crimpar terminal pino 0,25A 6mm Yyt-1', 'Ferramental', 3.37),
    ('Alicate para rebitar', 'Ferramental', 0.81),
    ('Alicate de crimpar conector BNC/Rg Yac-3Gs', 'Ferramental', 5.00),
    ('Alicate de crimpar plug Rg9/11/45 Ht-568R', 'Ferramental', 5.00),
    ('Decapador universal cabo coaxial Rg6/58/59', 'Ferramental', 1.44),
    ('Alicate universal isolado 8"', 'Ferramental', 1.45),
    ('Jogo de chave torx tipo L T10 a T50 9Pcs', 'Ferramental', 0.89),
    ('Chave de fenda com isolamento 1/8x6', 'Ferramental', 0.31),
    ('Chave de fenda com isolamento 6/16x6', 'Ferramental', 0.51),
    ('Chave de fenda com isolamento 1/4x6', 'Ferramental', 0.51),
    ('Philips 3/16 1-1/2 Philips 1/8x4 Philips 3/16x3', 'Ferramental', 4.76),
    ('Maleta de ferramentas', 'Ferramental', 6.25),
    ('Faca desencapadora alfa', 'Ferramental', 3.56),
    ('Ferro de solda 127V 40w Sc40Pd', 'Ferramental', 1.11),
    ('Jogo de chave allen 1,5 a 6mm 7 pecas', 'Ferramental', 1.56),
    ('Pincel 1.16/2"-38mm', 'Ferramental', 0.08),
    ('Nivel aluminio Gp12"', 'Ferramental', 0.72),
    ('Trena de aco L516CME-5M LUFKIN', 'Ferramental', 0.65),
    ('Chave canhao 8x245mm Belzer', 'Ferramental', 0.63),
    ('Chave canhao 10x245 Belzer', 'Ferramental', 0.83),
    ('Chave canhao 11x245 Belzer', 'Ferramental', 0.75),
    ('Chave canhao 12x245 Belzer', 'Ferramental', 0.90),
    ('Guia passa cabo', 'Ferramental', 0.83),
    ('Trena', 'Ferramental', 0.65),
    ('Arco de serra', 'Ferramental', 0.77),
    ('Jogo de chave catraca Bit para parafusadeira', 'Ferramental', 6.94),
    ('Escada de fibra tesoura de 04 degraus', 'Ferramental', 12.33),
    ('Escada de fibra tesoura de 08 degraus', 'Ferramental', 16.89),
    ('Escada de fibra extensivel 16/32DEG-9,9M', 'Ferramental', 44.42),
    ('Cones de seguranca pequeno', 'EPC', 0.41),
    ('Pedestais para corrente de seguranca', 'EPC', 0.83),
    ('Correntes de seguranca plastica amarela 10m', 'EPC', 0.11),
]

KIT_INSUMOS_ITEMS = [
    ('Fita Isolante 19MM x 20 metros 3M Scotch', 'Insumos', 33.00),
    ('Broca de aco rapido 3,5 MM', 'Insumos', 2.50),
    ('Broca de aco rapido 4,5 MM', 'Insumos', 2.50),
    ('Broca de videa 10MM', 'Insumos', 3.33),
    ('Broca de videa 12MM', 'Insumos', 4.17),
    ('Broca de videa 6MM', 'Insumos', 1.67),
    ('Broca de videa 8MM', 'Insumos', 2.50),
    ('Lamina serra bimetal starrett', 'Insumos', 20.00),
    ('Serra copa completo 20 25 e 35MM', 'Insumos', 22.50),
    ('Abracadeira nylon Hellermann T18 R', 'Insumos', 27.00),
    ('Abracadeira nylon Hellermann T30 L', 'Insumos', 26.00),
    ('Abracadeira nylon Hellermann T30 R', 'Insumos', 28.00),
    ('WD-40 300ml oleo anticorrosivo', 'Insumos', 40.00),
    ('Fita dupla face 3M ou Vonder', 'Insumos', 5.50),
    ('Parabolt 3/8', 'Insumos', 18.00),
    ('Parabolt 5/16', 'Insumos', 12.00),
    ('Trava rosca Loctite 242', 'Insumos', 20.00),
    ('Alcool Isopropilico', 'Insumos', 9.17),
    ('Tubo de silicone Sikaflex 295UV', 'Insumos', 30.00),
    ('Graxa Branca', 'Insumos', 7.33),
    ('Limpa contato componentes eletronicos Wurtl', 'Insumos', 10.67),
    ('Tubo de solda com fluxo', 'Insumos', 9.15),
]

KIT_ESCRITORIO_ITEMS = [
    ('Mesa para Containers Adm e Almoxarifado', 'Escritorio', 16.67),
    ('Cadeiras Adm e Almoxarifado', 'Escritorio', 2.78),
    ('Impressora Multifuncional Adm', 'Escritorio', 22.22),
    ('Insumos para Impressora', 'Consumiveis', 25.00),
    ('Armarios Adm', 'Escritorio', 13.89),
    ('Limpeza dos 2 Containers Adm e Almox', 'Consumiveis', 8.33),
    ('Material Escritorio', 'Consumiveis', 83.33),
]

LISTA_ATIVOS_FERR = [
    # (descricao, unidade) — sem preço fixo; default vem do histórico de orçamentos
    ('Mobilização',                               'Vb'),
    ('Desmobilização',                            'Vb'),
    ('Container',                                 'Meses'),
    ('Locação de Rádio Comunicador',              'Meses'),
    ('Locação de Kit de Fusão Óptica',            'Meses'),
    ('Locação de Certificadora de Cabo Metálico', 'Meses'),
]

KITS_FERR = {
    'kit_maleta': {
        'label': 'Maleta de Ferramentas',
        'desc': 'Alocação por equipe operativa',
        'icon': 'home_repair_service',
        'cor_borda': 'border-primary',
        'cor_label': 'text-primary',
        'multiplicador': 'profissional',
        'items': KIT_MALETA_ITEMS,
    },
    'kit_insumos': {
        'label': 'Insumos Perecíveis',
        'desc': 'Gestão de estoque crítico',
        'icon': 'inventory_2',
        'cor_borda': 'border-amber-500',
        'cor_label': 'text-amber-600',
        'multiplicador': 'profissional',
        'items': KIT_INSUMOS_ITEMS,
    },
    'kit_escritorio': {
        'label': 'Escritório de Obra',
        'desc': 'Custo fixo de manutenção',
        'icon': 'domain',
        'cor_borda': 'border-teal-500',
        'cor_label': 'text-teal-600',
        'multiplicador': 'fixo',
        'items': KIT_ESCRITORIO_ITEMS,
    },
}

# Funções operacionais que recebem kit por pessoa (Dupla Técnica = 2 por alocação)
MULT_KIT_PROFISSIONAL = {
    'Dupla Técnica': 2, 'Dupla Técnica (Civil)': 2,
    'Auxiliar de Instalação': 1, 'Eletricista': 1,
    'Técnico de Manutenção': 1, 'Técnico de TI': 1, 'Técnico de T.I': 1,
    'Técnico de TI - Ferroport': 1, 'Técnico de Manutenção - Ferroport': 1,
    'Técnico de Manutenção - Prumo': 1, 'Técnico de Segurança do Trabalho': 1,
    'Técnico de Manutenção - Samarco': 1, 'Pedreiro': 1, 'Ajudante de Pedreiro': 1,
}

# (slug, label, icon, custo, unidade)
# Fusão e Certificação migraram para a Lista de Ativos (preço editável via histórico)
EQUIPS_FERR = [
    ('caminhao_plataforma', 'Caminhão Plataforma', 'local_shipping',   1400, 'Dias'),
    ('caminhao_munk',       'Caminhão Munk',        'construction',    1350, 'Dias'),
    ('pta_elevatoria',      'PTA Plat. Elevatória', 'elevator',        1500, 'Dias'),
    ('exames_treinamento',  'Exames e Treinamento', 'health_and_safety',  0, 'Vb'),
]


def _headcount_demais_custos(projeto):
    """Conta cabeças para benefícios: Dupla Técnica = 2 pessoas, demais = 1 por unidade."""
    itens_op = ItemMO.objects.filter(projeto=projeto, aba='operacional', ativo=True)
    total = 0
    for item in itens_op:
        mult = 2 if item.descricao in ('Dupla Técnica', 'Dupla Técnica (Civil)') else 1
        total += int(item.quantidade) * mult
    return max(total, 1)


def _ultimo_preco_ativo(descricao, projeto_atual):
    """Retorna o último custo_unitario para um item da Lista de Ativos, priorizando
    mesmo município → mesmo estado → qualquer projeto. Exclui o projeto atual."""
    from decimal import Decimal
    base_qs = (
        ItemMO.objects
        .filter(aba='ferramentas', descricao=descricao, especificacao='')
        .exclude(projeto=projeto_atual)
        .select_related('projeto')
        .order_by('-projeto__data_criacao')
    )
    for filtro in [
        {'projeto__municipio_obra': projeto_atual.municipio_obra,
         'projeto__estado_obra':    projeto_atual.estado_obra},
        {'projeto__estado_obra': projeto_atual.estado_obra},
        {},
    ]:
        hit = base_qs.filter(**filtro).first()
        if hit:
            return hit.custo_unitario
    return Decimal('0')


def _kit_qtd_default(projeto, multiplicador_tipo):
    """Calcula quantidade padrão de kits baseada na equipe operacional."""
    from decimal import Decimal
    if multiplicador_tipo == 'fixo':
        return Decimal('1')
    itens_op = ItemMO.objects.filter(projeto=projeto, aba='operacional', ativo=True)
    total = 0
    for item in itens_op:
        mult = MULT_KIT_PROFISSIONAL.get(item.descricao, 0)
        total += int(item.quantidade) * mult
    return Decimal(str(max(total, 1)))


@login_required(login_url='/')
def gestao_mo(request, projeto_id, aba):
    from decimal import Decimal, InvalidOperation

    if aba not in [a[0] for a in ABAS_MO]:
        return redirect('gestao_mo', projeto_id=projeto_id, aba='operacional')
    projeto = get_object_or_404(Projeto, pk=projeto_id)

    # ── Métricas da equipe operacional (usadas em todas as abas) ────────────
    import math as _math
    _ops_ativos = ItemMO.objects.filter(projeto=projeto, aba='operacional', ativo=True)
    _meses_list = [i.tempo for i in _ops_ativos if i.unidade == 'Meses']
    equipe_meses = max(_meses_list) if _meses_list else Decimal('1')
    equipe_headcount = _headcount_demais_custos(projeto)
    equipe_dias_uteis = int(equipe_meses) * 22
    veiculos_sugeridos = _math.ceil(equipe_headcount / 4) if equipe_headcount > 0 else 1

    def _dec(val, default=0):
        try:
            return Decimal(str(val).replace(',', '.'))
        except (InvalidOperation, TypeError):
            return Decimal(str(default))

    # ── Auto-populate demais_custos on first visit ──────────────────────────
    if aba == 'demais_custos' and not ItemMO.objects.filter(projeto=projeto, aba='demais_custos').exists():
        headcount_dc = _headcount_demais_custos(projeto)
        tem_viagem = ItemMO.objects.filter(
            projeto=projeto, aba='transporte', descricao='Combustível - Viagem Longa', ativo=True
        ).exists()
        VIAGEM_ITENS = {'Alimentação (Diária)', 'Hospedagem'}
        for desc, _, custo, unidade, tempo, ativo_def in ITENS_DEMAIS_CUSTOS_DEFAULT:
            ativo_real = ativo_def if desc not in VIAGEM_ITENS else (ativo_def and tem_viagem)
            ItemMO.objects.create(
                projeto=projeto, aba='demais_custos',
                descricao=desc, quantidade=headcount_dc, tempo=tempo,
                unidade=unidade, custo_unitario=Decimal(str(custo)), ativo=ativo_real,
            )

    # ── Auto-populate ferramentas lista de ativos on first visit ────────────
    if aba == 'ferramentas' and not ItemMO.objects.filter(
        projeto=projeto, aba='ferramentas', especificacao=''
    ).exists():
        for nome, unidade in LISTA_ATIVOS_FERR:
            preco = _ultimo_preco_ativo(nome, projeto)
            ItemMO.objects.create(
                projeto=projeto, aba='ferramentas',
                descricao=nome, especificacao='',
                quantidade=Decimal('1'), tempo=Decimal('1'),
                unidade=unidade, custo_unitario=preco, ativo=False,
            )

    # ── Auto-populate terceiros on first visit ───────────────────────────────
    if aba == 'terceiros' and not ItemMO.objects.filter(projeto=projeto, aba='terceiros').exists():
        for desc, espec, custo in ITENS_TERCEIROS_DEFAULT:
            ItemMO.objects.create(
                projeto=projeto, aba='terceiros',
                descricao=desc, especificacao=espec,
                quantidade=0, tempo=1, unidade='Un',
                custo_unitario=Decimal(str(custo)), ativo=False, status='pendente',
            )

    if request.method == 'POST':
        acao = request.POST.get('acao')

        if acao == 'add':
            descricao = request.POST.get('descricao', '').strip()
            if descricao:
                ItemMO.objects.create(
                    projeto=projeto, aba=aba, descricao=descricao,
                    especificacao=request.POST.get('especificacao', '').strip(),
                    quantidade=_dec(request.POST.get('quantidade', 1), 1),
                    tempo=_dec(request.POST.get('tempo', 1), 1),
                    unidade=request.POST.get('unidade', 'Meses'),
                    custo_unitario=_dec(request.POST.get('custo_unitario', 0)),
                    ativo=True,
                )

        elif acao == 'delete':
            ItemMO.objects.filter(pk=request.POST.get('item_id'), projeto=projeto).delete()

        elif acao == 'update':
            item = get_object_or_404(ItemMO, pk=request.POST.get('item_id'), projeto=projeto)
            new_unit = request.POST.get('unidade', item.unidade)
            new_custo = _dec(request.POST.get('custo_unitario', item.custo_unitario), item.custo_unitario)
            # Se mudou unidade e função tem salário base conhecido, recalcula automaticamente
            if aba == 'operacional' and new_unit != item.unidade and item.descricao in SALARIOS_MO:
                base = Decimal(str(SALARIOS_MO[item.descricao]))
                if new_unit == 'Meses':
                    new_custo = base
                elif new_unit == 'Dias':
                    new_custo = (base / 20).quantize(Decimal('0.01'))
                elif new_unit == 'Horas':
                    new_custo = (base / 146).quantize(Decimal('0.01'))
            item.quantidade = _dec(request.POST.get('quantidade', item.quantidade), item.quantidade)
            item.tempo = _dec(request.POST.get('tempo', item.tempo), item.tempo)
            item.unidade = new_unit
            item.custo_unitario = new_custo
            item.especificacao = request.POST.get('especificacao', item.especificacao)
            item.status = request.POST.get('status', item.status)
            item.ativo = request.POST.get('ativo') == '1'
            item.save()

        elif acao == 'toggle_ativo':
            item = get_object_or_404(ItemMO, pk=request.POST.get('item_id'), projeto=projeto)
            item.ativo = not item.ativo
            if item.ativo and item.status == 'pendente':
                item.status = 'ativo'
            elif not item.ativo:
                item.status = 'pendente'
            item.save()

        elif acao == 'frota_update':
            # Transporte: salva/atualiza cada veículo da frota
            # template usa {{ label|cut:" " }} → removes ALL spaces, no underscore
            for label, icon, custo_base in FROTA_TRANSPORTE:
                field_key = 'frota_' + label.replace(' ', '')
                qtd_raw = request.POST.get(field_key, '0')
                qtd = _dec(qtd_raw, 0)
                custo_base_dec = Decimal(str(custo_base))
                if qtd > 0:
                    obj, _ = ItemMO.objects.get_or_create(
                        projeto=projeto, aba='transporte', descricao=label,
                        defaults={'custo_unitario': custo_base_dec, 'tempo': 1, 'unidade': 'Meses'},
                    )
                    obj.quantidade = qtd
                    obj.custo_unitario = custo_base_dec
                    obj.ativo = True
                    obj.save()
                else:
                    ItemMO.objects.filter(projeto=projeto, aba='transporte', descricao=label).delete()
            # Viagem longa e deslocamento diário ficam como items separados
            for campo, desc in [('km_viagem', 'Combustível - Viagem Longa'), ('km_diario', 'Combustível - Deslocamento Diário')]:
                val = _dec(request.POST.get(campo, 0))
                if val > 0:
                    ItemMO.objects.update_or_create(
                        projeto=projeto, aba='transporte', descricao=desc,
                        defaults={'quantidade': 1, 'tempo': 1, 'unidade': 'Vb', 'custo_unitario': val, 'ativo': True},
                    )

        elif acao == 'kit_toggle':
            kit_key = request.POST.get('kit_key', '')
            if kit_key in KITS_FERR:
                kit_def = KITS_FERR[kit_key]
                has_active = ItemMO.objects.filter(
                    projeto=projeto, aba='ferramentas', especificacao=kit_key, ativo=True
                ).exists()
                if has_active:
                    # Desativar todos os itens do kit
                    ItemMO.objects.filter(
                        projeto=projeto, aba='ferramentas', especificacao=kit_key
                    ).update(ativo=False)
                else:
                    # Ativar: criar/reativar itens com qtd baseada na equipe
                    qtd_default = _kit_qtd_default(projeto, kit_def['multiplicador'])
                    # Remove os inativos e recria
                    ItemMO.objects.filter(
                        projeto=projeto, aba='ferramentas', especificacao=kit_key
                    ).delete()
                    for nome, grupo, custo_base in kit_def['items']:
                        ItemMO.objects.create(
                            projeto=projeto, aba='ferramentas', especificacao=kit_key,
                            descricao=nome,
                            quantidade=qtd_default, tempo=equipe_meses, unidade='Meses',
                            custo_unitario=Decimal(str(custo_base)), ativo=True,
                        )

        elif acao == 'kit_save':
            kit_key = request.POST.get('kit_key', '')
            if kit_key in KITS_FERR:
                kit_def = KITS_FERR[kit_key]
                kit_unidade = 'Meses'
                kit_periodo = equipe_meses  # período vem da equipe operacional
                # Apaga e recria para evitar duplicatas
                ItemMO.objects.filter(
                    projeto=projeto, aba='ferramentas', especificacao=kit_key
                ).delete()
                for i, (nome, grupo, custo_base) in enumerate(kit_def['items']):
                    ativo = request.POST.get(f'item_ativo_{i}') == '1'
                    qty = _dec(request.POST.get(f'item_qty_{i}', '1'), 1)
                    monthly = Decimal(str(custo_base))
                    if kit_unidade == 'Dias':
                        custo_unit = (monthly / 20).quantize(Decimal('0.01'))
                    else:
                        custo_unit = monthly
                    ItemMO.objects.create(
                        projeto=projeto, aba='ferramentas', especificacao=kit_key,
                        descricao=nome,
                        quantidade=qty, tempo=kit_periodo, unidade=kit_unidade,
                        custo_unitario=custo_unit, ativo=ativo,
                    )

        elif acao == 'equip_toggle':
            equip_slug = request.POST.get('equip', '')
            equip_def = next((e for e in EQUIPS_FERR if e[0] == equip_slug), None)
            if equip_def:
                _, label, _, custo, unidade = equip_def
                existing = ItemMO.objects.filter(
                    projeto=projeto, aba='ferramentas', especificacao='equip', descricao=label
                ).first()
                if existing:
                    existing.ativo = not existing.ativo
                    existing.save(update_fields=['ativo'])
                else:
                    ItemMO.objects.create(
                        projeto=projeto, aba='ferramentas', especificacao='equip',
                        descricao=label, quantidade=Decimal('1'), tempo=Decimal('1'),
                        unidade=unidade, custo_unitario=Decimal(str(custo)), ativo=True,
                    )

        elif acao == 'aplicar_bom':
            total_mo = sum(
                i.custo_total for i in ItemMO.objects.filter(projeto=projeto, ativo=True)
            )
            for item_proj in ItemProjeto.objects.filter(projeto=projeto, produto__preco_variavel=True):
                qtd = item_proj.quantidade or 1
                item_proj.preco_unitario = (total_mo / qtd).quantize(Decimal('0.01'))
                item_proj.save(update_fields=['preco_unitario'])
            return redirect('bom_selector_projeto', projeto_id=projeto_id)

        return redirect('gestao_mo', projeto_id=projeto_id, aba=aba)

    # ── Context ──────────────────────────────────────────────────────────────
    itens = list(ItemMO.objects.filter(projeto=projeto, aba=aba))
    total_aba = sum(i.custo_total for i in itens if i.ativo)
    total_mo = sum(
        i.custo_total for i in ItemMO.objects.filter(projeto=projeto, ativo=True)
    )

    # Total de horas (operacional)
    total_horas = Decimal('0')
    if aba == 'operacional':
        for i in itens:
            if i.unidade == 'Horas':
                total_horas += i.quantidade * i.tempo
            elif i.unidade == 'Dias':
                total_horas += i.quantidade * i.tempo * 8
            elif i.unidade == 'Meses':
                total_horas += i.quantidade * i.tempo * 176

    # % M.O no projeto
    total_bom = sum(
        ip.preco_unitario * ip.quantidade
        for ip in ItemProjeto.objects.filter(projeto=projeto)
    )
    pct_mo = round((total_mo / total_bom * 100), 1) if total_bom > 0 else 0

    # Frota transporte: lista de (label, icon, custo_base, qtd_atual)
    frota_lista = []
    comb_viagem_custo = 0.0
    comb_diario_custo = 0.0
    if aba == 'transporte':
        for label, icon, custo_base in FROTA_TRANSPORTE:
            item_fr = ItemMO.objects.filter(projeto=projeto, aba='transporte', descricao=label).first()
            qtd_atual = int(item_fr.quantidade) if item_fr else 0
            frota_lista.append((label, icon, custo_base, qtd_atual))
        comb_v = ItemMO.objects.filter(projeto=projeto, aba='transporte', descricao='Combustível - Viagem Longa').first()
        comb_d = ItemMO.objects.filter(projeto=projeto, aba='transporte', descricao='Combustível - Deslocamento Diário').first()
        comb_viagem_custo = float(comb_v.custo_unitario) if comb_v else 0.0
        comb_diario_custo = float(comb_d.custo_unitario) if comb_d else 0.0

    # ── Kits e Equipamentos (só para aba ferramentas) ────────────────────────
    kits_data = {}
    equips_data = []
    if aba == 'ferramentas':
        qtd_prof = _kit_qtd_default(projeto, 'profissional')
        qtd_fixo = Decimal('1')
        for kit_key, kit_def in KITS_FERR.items():
            existing_map = {
                item.descricao: item
                for item in ItemMO.objects.filter(
                    projeto=projeto, aba='ferramentas', especificacao=kit_key
                )
            }
            mult_tipo = kit_def['multiplicador']
            qtd_default = qtd_prof if mult_tipo == 'profissional' else qtd_fixo
            kit_items_data = []
            kit_total = Decimal('0')
            has_active = False
            for i, (nome, grupo, custo_base) in enumerate(kit_def['items']):
                db = existing_map.get(nome)
                qty = db.quantidade if db else qtd_default
                ativo_default_kit = (kit_def['multiplicador'] != 'fixo')  # escritório=False, outros=True
                ativo = db.ativo if db else ativo_default_kit
                custo = Decimal(str(custo_base))
                item_total = qty * custo if ativo else Decimal('0')
                kit_total += item_total
                if ativo:
                    has_active = True
                kit_items_data.append({
                    'idx': i, 'nome': nome, 'grupo': grupo,
                    'custo_base': float(custo_base),
                    'qty': float(qty), 'ativo': ativo, 'total': float(item_total),
                })
            kits_data[kit_key] = {
                'key': kit_key,
                'label': kit_def['label'], 'desc': kit_def['desc'],
                'icon': kit_def['icon'],
                'cor_borda': kit_def['cor_borda'], 'cor_label': kit_def['cor_label'],
                'items_data': kit_items_data,
                'total': kit_total,
                'is_active': has_active,
            }
        for slug, label, icon, custo, unidade in EQUIPS_FERR:
            db = ItemMO.objects.filter(
                projeto=projeto, aba='ferramentas', especificacao='equip', descricao=label
            ).first()
            equips_data.append({
                'slug': slug, 'label': label, 'icon': icon,
                'custo': custo, 'unidade': unidade,
                'is_active': db.ativo if db else False,
                'item_id': db.id if db else None,
            })

    import json as _json
    # Operacional: lista de (nome, salario_mensal) para auto-fill de custo no JS
    funcoes_op = [
        (nome, SALARIOS_MO.get(nome, 0))
        for nome in FUNCOES_MO.get('operacional', [])
    ]
    salarios_json = _json.dumps(SALARIOS_MO)
    funcao_icons_json = _json.dumps(FUNCAO_ICONS)

    encargos_valor = (total_aba * Decimal(str(ENCARGOS_PERCENT)) / 100).quantize(Decimal('0.01'))
    total_com_encargos = total_aba + encargos_valor

    # ── Auto-sync item preco_variavel: sempre reflete o total M.O ───────────
    _prod_var = Produto.objects.filter(preco_variavel=True).first()
    if _prod_var:
        ItemProjeto.objects.update_or_create(
            projeto=projeto, produto=_prod_var,
            defaults={
                'quantidade': 1,
                'preco_unitario': total_mo,
                'faturar_servico': True,
            }
        )

    total_itens_projeto = ItemProjeto.objects.filter(projeto=projeto).count()

    ctx = {
        'projeto': projeto,
        'aba': aba,
        'abas_mo': ABAS_MO,
        'total_itens_projeto': total_itens_projeto,
        'itens': itens,
        'total_aba': total_aba,
        'encargos_percent': ENCARGOS_PERCENT,
        'encargos_valor': encargos_valor,
        'total_com_encargos': total_com_encargos,
        'total_mo': total_mo,
        'total_horas': int(total_horas),
        'pct_mo': pct_mo,
        'total_bom': total_bom,
        'funcoes_aba': FUNCOES_MO.get(aba, []),   # strings (ferramentas/transporte/etc)
        'funcoes_op': funcoes_op,                  # tuplas (nome, salario) só p/ operacional
        'salarios_json': salarios_json,
        'funcao_icons_json': funcao_icons_json,
        'cfg': CONFIG_ABA.get(aba, CONFIG_ABA['operacional']),
        'frota_lista': frota_lista,
        'comb_viagem_custo': comb_viagem_custo,
        'comb_diario_custo': comb_diario_custo,
        'kits_data': kits_data,
        'equips_data': equips_data,
        'kits_ferr_json': _json.dumps({k: v['label'] for k, v in KITS_FERR.items()}),
        'lista_ativos_nomes': [n for n, _ in LISTA_ATIVOS_FERR],
        'equipe_meses': int(equipe_meses),
        'equipe_headcount': equipe_headcount,
        'equipe_dias_uteis': equipe_dias_uteis,
        'veiculos_sugeridos': veiculos_sugeridos,
    }
    return render(request, 'gestao_mo.html', ctx)


# ============================================================
# DASHBOARDS
# ============================================================
@login_required(login_url='/')
def dash_gerencial(request):
    cargo = get_cargo(request.user)
    # Cargos sem acesso ao dash gerencial são redirecionados ao seu dash
    if cargo in CARGOS_COMPRAS:
        return redirect('dash_orcamentista')
    if cargo == 'analista':
        return redirect('dash_analista')
    return render(request, 'dash_Gerencial.html')

@login_required(login_url='/')
def dash_analista(request):
    cargo = get_cargo(request.user)
    if cargo in CARGOS_COMPRAS:
        return redirect('dash_orcamentista')
    if cargo not in {'analista'} | CARGOS_DIRETORIA | CARGOS_GERENCIA:
        return redirect('dash_gerencial')
    return render(request, 'dash_analista.html')

@login_required(login_url='/')
def dash_orcamentista(request):
    cargo = get_cargo(request.user)
    if cargo not in CARGOS_COMPRAS | CARGOS_DIRETORIA | CARGOS_GERENCIA:
        return redirect('dash_analista')
    return render(request, 'dash_orçamentista.html')

@login_required(login_url='/')
def validacao_orcamento(request):
    cargo = get_cargo(request.user)
    if cargo not in CARGOS_DIRETORIA | CARGOS_GERENCIA | CARGOS_COMPRAS:
        messages.error(request, 'Você não tem permissão para acessar esta página.')
        return redirect('inicio')
    return render(request, 'validacao_de_orcamento.html')


# ============================================================
# DIAGNÓSTICO (sem login — apenas para verificar o sistema)
# ============================================================
from django.http import JsonResponse

def status_diagnostico(request):
    from django.conf import settings
    import os, json, tempfile
    from django.core.management import call_command

    fixture_path = os.path.join(settings.BASE_DIR, 'fixtures', 'initial_data.json')
    fixture_existe = os.path.exists(fixture_path)

    # Upload de logos via ?upload=1
    upload_resultado = None
    if request.GET.get('upload') == '1':
        try:
            from django.core.management import call_command as _cc
            _cc('upload_media_to_gcs', verbosity=0)
            upload_resultado = 'OK'
        except Exception as e:
            upload_resultado = f'ERRO: {e}'

    # Carga de produtos via ?load_products=1
    carga_produtos_resultado = None
    if request.GET.get('load_products') == '1' and fixture_existe:
        try:
            if Produto.objects.exists():
                carga_produtos_resultado = f'Produtos já existem ({Produto.objects.count()}), ignorado'
            else:
                # Carrega apenas os modelos que ainda não existem
                ONLY_PRODUCTS = {'core.produto', 'core.projeto', 'core.itemprojeto'}
                with open(fixture_path, encoding='utf-8') as f:
                    all_data = json.load(f)
                products_data = [o for o in all_data if o['model'] in ONLY_PRODUCTS]
                with tempfile.NamedTemporaryFile(mode='w', suffix='.json',
                                                 delete=False, encoding='utf-8') as tmp:
                    json.dump(products_data, tmp, ensure_ascii=False)
                    tmp_path = tmp.name
                try:
                    call_command('loaddata', tmp_path, verbosity=0)
                    carga_produtos_resultado = f'OK — {Produto.objects.count()} produtos carregados'
                finally:
                    os.unlink(tmp_path)
        except Exception as e:
            carga_produtos_resultado = f'ERRO: {e}'

    # Carga manual via ?load=1
    carga_resultado = None
    if request.GET.get('load') == '1' and fixture_existe:
        try:
            if not Empresa.objects.exists():
                with open(fixture_path, encoding='utf-8') as f:
                    all_data = json.load(f)
                priority_data = [o for o in all_data
                                 if o['model'] in ('auth.group', 'core.empresa', 'auth.user', 'core.perfil')]
                with tempfile.NamedTemporaryFile(mode='w', suffix='.json',
                                                 delete=False, encoding='utf-8') as tmp:
                    json.dump(priority_data, tmp, ensure_ascii=False)
                    tmp_path = tmp.name
                try:
                    call_command('loaddata', tmp_path, verbosity=0)
                    carga_resultado = f'OK — {len(priority_data)} registros carregados'
                finally:
                    os.unlink(tmp_path)
            else:
                carga_resultado = 'Empresas já existem, carga ignorada'
        except Exception as e:
            carga_resultado = f'ERRO: {e}'

    # Contagem do banco
    try:
        db_empresas = Empresa.objects.count()
        db_usuarios = User.objects.count()
        db_produtos = Produto.objects.count()
        db_projetos = Projeto.objects.count()
        db_ok = True
        db_erro = None
    except Exception as e:
        db_ok = False
        db_erro = str(e)
        db_empresas = db_usuarios = db_produtos = db_projetos = 0

    empresas_lista = []
    if db_ok:
        for e in Empresa.objects.all():
            empresas_lista.append({
                'id': e.id,
                'nome': e.nome_fantasia,
                'logo': str(e.logo) if e.logo else None,
                'is_sistema': e.is_sistema,
            })

    # GCS
    gcs_ok = False
    gcs_erro = None
    if getattr(settings, 'USE_GCS', False):
        try:
            from google.cloud import storage as gcs
            client = gcs.Client()
            bucket = client.bucket(settings.GS_BUCKET_NAME)
            gcs_ok = bucket.exists()
        except Exception as e:
            gcs_erro = str(e)

    return JsonResponse({
        'banco': {
            'ok': db_ok,
            'erro': db_erro,
            'empresas': db_empresas,
            'usuarios': db_usuarios,
            'produtos': db_produtos,
            'projetos': db_projetos,
        },
        'empresas_cadastradas': empresas_lista,
        'fixture': {
            'path': fixture_path,
            'existe': fixture_existe,
        },
        'carga_manual': carga_resultado,
        'carga_produtos': carga_produtos_resultado,
        'upload_logos': upload_resultado,
        'gcs': {
            'ativo': getattr(settings, 'USE_GCS', False),
            'bucket': getattr(settings, 'GS_BUCKET_NAME', None),
            'bucket_existe': gcs_ok,
            'erro': gcs_erro,
        },
        'debug': settings.DEBUG,
    }, json_dumps_params={'ensure_ascii': False, 'indent': 2})
